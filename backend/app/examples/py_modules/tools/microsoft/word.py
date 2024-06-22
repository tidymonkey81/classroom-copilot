import os
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage
import io

def open_docx(docx_path):
    if os.path.exists(docx_path):
        document = Document(docx_path)
    else:
        document = Document()
    return document

# function to close a docx document
def close_docx(document):
    document.close()

def save_docx(document, docx_path):
    # Check if the directory exists and if not create it
    dir_path = os.path.dirname(docx_path)
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    document.save(docx_path)

def add_paragraph(document, text, style=None):
    paragraph = document.add_paragraph(text, style)
    return paragraph

def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    return heading

def add_image_old(image_path, document, width=None):
    if not os.path.exists(image_path):
        print("Image does not exist: " + image_path)
        return
    if width:
        width = Inches(width)
    picture = document.add_picture(image_path, width=width)
    return picture

def add_section(document, title=None, text=None, list=None):
    if title:
        add_heading(document, title, level=1)
    if text:
        add_paragraph(document, text)
    if list:
        add_list(document, list)

def add_list(document, list_items, style=None):
    if style:
        style = document.styles[style]
    else:
        style = document.styles['List Bullet']
    for item in list_items:
        print("Item: " + item)
        print("Adding paragraph...")
        add_paragraph(document, item, style=style)

def add_image(image_path, document, width=None):
    # Check if the image file exists
    if not os.path.exists(image_path):
        print(f"Image does not exist: {image_path}")
        return None

    try:
        # Ensure the image is in a compatible format
        with PILImage.open(image_path) as img:
            # Convert image to RGB if it's not
            if img.mode not in ['RGB', 'RGBA']:
                img = img.convert('RGB')

            # Save the image to a BytesIO object
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr = img_byte_arr.getvalue()

            # Add the image from BytesIO object to the document
            if width:
                width = Inches(width)
            picture = document.add_picture(io.BytesIO(img_byte_arr), width=width)
            return picture

    except Exception as e:
        print(f"Error adding image '{image_path}': {e}")
        return None